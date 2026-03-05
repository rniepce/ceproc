"""
DPT Generator — Gerador do Documento Descritivo do Processo de Trabalho.
========================================================================
Percorre o grafo BPMN parseado e gera um documento DOCX estruturado
com passo a passo narrativo, glossário de sistemas, insumos/produtos
e desvios lógicos.
"""

from collections import deque
from dataclasses import dataclass, field
from typing import Optional
import io
import re

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from bpmn_parser import BpmnModel, BpmnElement


# ── TJMG Brand Colors ──────────────────────────────────────────────
TJMG_VINHO = RGBColor(0x6B, 0x1D, 0x2A)
TJMG_AZUL = RGBColor(0x1B, 0x2A, 0x4A)
TJMG_DOURADO = RGBColor(0xC5, 0xA5, 0x5A)
TJMG_CINZA = RGBColor(0x66, 0x66, 0x66)


# ── Helper: Check if name is an abbreviation ───────────────────────
def _is_abbreviation(name: str) -> bool:
    """Check if a system name is an abbreviation (all caps, short)."""
    clean = name.strip()
    return (
        len(clean) <= 8 and
        clean == clean.upper() and
        any(c.isalpha() for c in clean)
    )


# ── Helper: Style a paragraph ─────────────────────────────────────
def _style_run(run, font_name="Segoe UI", font_size=10, color=None,
               bold=False, italic=False):
    """Apply styling to a docx Run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def _add_styled_paragraph(doc, text, font_size=10, color=None,
                           bold=False, italic=False, alignment=None,
                           space_after=6, space_before=0):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _style_run(run, font_size=font_size, color=color, bold=bold, italic=italic)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def _add_heading_styled(doc, text, level=1, color=None):
    """Add a heading with custom color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Segoe UI"
        if color:
            run.font.color.rgb = color
    return heading


# ── Build connected data associations ──────────────────────────────

def _build_associations_map(model: BpmnModel):
    """
    Build maps:
    - task_annotations: task_id → [annotation_text, ...]
    - task_data_in: task_id → [data_object_name, ...]  (insumos)
    - task_data_out: task_id → [data_object_name, ...]  (produtos)
    - task_data_store: task_id → data_store_name (connected system)
    """
    task_annotations = {}
    task_data_in = {}
    task_data_out = {}
    task_data_store = {}

    # Process associations (annotation ↔ element)
    for assoc in model.associations:
        src = assoc.source_ref
        tgt = assoc.target_ref

        # Annotation → Element
        if src in model.annotations:
            ann_text = model.annotations[src].name
            task_annotations.setdefault(tgt, []).append(ann_text)
        elif tgt in model.annotations:
            ann_text = model.annotations[tgt].name
            task_annotations.setdefault(src, []).append(ann_text)

    # Process data object/store connections via sequence flows
    # and data associations within the XML
    for flow_id, flow in model.flows.items():
        src = flow.source_ref
        tgt = flow.target_ref

        # DataObject → Task (insumo)
        if src in model.data_objects and tgt in model.elements:
            obj = model.data_objects[src]
            if obj.data_type == "dataStore":
                task_data_store[tgt] = obj.name
            else:
                task_data_in.setdefault(tgt, []).append(obj.name)

        # Task → DataObject (produto)
        if src in model.elements and tgt in model.data_objects:
            obj = model.data_objects[tgt]
            if obj.data_type == "dataStore":
                task_data_store[src] = obj.name
            else:
                task_data_out.setdefault(src, []).append(obj.name)

    # Also check associations for data objects
    for assoc in model.associations:
        src = assoc.source_ref
        tgt = assoc.target_ref

        if src in model.data_objects:
            obj = model.data_objects[src]
            if obj.data_type == "dataStore":
                task_data_store[tgt] = obj.name
            else:
                task_data_in.setdefault(tgt, []).append(obj.name)

        if tgt in model.data_objects:
            obj = model.data_objects[tgt]
            if obj.data_type == "dataStore":
                task_data_store[src] = obj.name
            else:
                task_data_out.setdefault(src, []).append(obj.name)

    return task_annotations, task_data_in, task_data_out, task_data_store


# ═══════════════════════════════════════════════════════════════════
# MAIN DPT GENERATOR
# ═══════════════════════════════════════════════════════════════════

def generate_dpt(model: BpmnModel) -> bytes:
    """
    Generate a DPT (Documento Descritivo do Processo de Trabalho) as DOCX.

    Traverses the BPMN graph from start events, building a narrative
    document with actors, systems glossary, step-by-step workflow,
    data objects, annotations, and gateway decisions.

    Returns: DOCX file as bytes.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(10)

    # Build association maps
    task_annotations, task_data_in, task_data_out, task_data_store = \
        _build_associations_map(model)

    # ── CAPA ────────────────────────────────────────────────────────
    _add_styled_paragraph(
        doc, "TRIBUNAL DE JUSTIÇA DO ESTADO DE MINAS GERAIS",
        font_size=16, color=TJMG_VINHO, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    _add_styled_paragraph(
        doc, "CEPROC — Centro de Estudos de Procedimentos",
        font_size=11, color=TJMG_AZUL,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    # Divider line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("━" * 60)
    _style_run(run, font_size=8, color=TJMG_DOURADO)

    _add_styled_paragraph(
        doc, "DOCUMENTO DESCRITIVO DO PROCESSO DE TRABALHO (DPT)",
        font_size=14, color=TJMG_AZUL, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18
    )

    # Header data
    if model.header:
        h = model.header
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers_data = [
            ("Título:", h.title or model.process_name or "—"),
            ("Autor:", h.author or "CEPROC"),
            ("Versão:", h.version or "—"),
            ("Descrição:", h.description or "—"),
        ]

        for i, (label, value) in enumerate(headers_data):
            cell_label = table.cell(i, 0)
            cell_value = table.cell(i, 1)

            p_label = cell_label.paragraphs[0]
            run_label = p_label.add_run(label)
            _style_run(run_label, font_size=10, color=TJMG_VINHO, bold=True)

            p_value = cell_value.paragraphs[0]
            run_value = p_value.add_run(value)
            _style_run(run_value, font_size=10)

        # Style table
        for row in table.rows:
            for cell in row.cells:
                cell.width = Cm(8.5)
    else:
        _add_styled_paragraph(
            doc, model.process_name or "Processo",
            font_size=14, color=TJMG_VINHO, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

    doc.add_page_break()

    # ── ATORES ──────────────────────────────────────────────────────
    _add_heading_styled(doc, "1. Atores do Processo", level=1, color=TJMG_VINHO)

    if model.lanes:
        table = doc.add_table(rows=1 + len(model.lanes), cols=2)
        # Header row
        h_cells = table.rows[0].cells
        for cell, text in zip(h_cells, ["#", "Ator / Setor"]):
            p = cell.paragraphs[0]
            run = p.add_run(text)
            _style_run(run, font_size=9, color=TJMG_AZUL, bold=True)

        for i, lane in enumerate(model.lanes):
            row = table.rows[i + 1]
            p_num = row.cells[0].paragraphs[0]
            run_num = p_num.add_run(str(i + 1))
            _style_run(run_num, font_size=9)

            p_name = row.cells[1].paragraphs[0]
            run_name = p_name.add_run(lane.name)
            _style_run(run_name, font_size=9)
    else:
        _add_styled_paragraph(doc, "Nenhum ator identificado (sem lanes definidas).",
                               font_size=10, italic=True, color=TJMG_CINZA)

    # ── GLOSSÁRIO DE SISTEMAS ───────────────────────────────────────
    data_stores = [obj for obj in model.data_objects.values()
                   if obj.data_type == "dataStore"]

    if data_stores:
        _add_heading_styled(doc, "2. Glossário de Sistemas", level=1, color=TJMG_VINHO)

        for ds in data_stores:
            p = doc.add_paragraph()
            run_name = p.add_run(f"• {ds.name}")
            _style_run(run_name, font_size=10, bold=True)

            if _is_abbreviation(ds.name):
                p_desc = doc.add_paragraph()
                run_desc = p_desc.add_run(
                    f"  [Preencher a descrição completa da sigla: {ds.name}]"
                )
                _style_run(run_desc, font_size=9, italic=True, color=TJMG_CINZA)

    # ── PASSO A PASSO DO PROCESSO ───────────────────────────────────
    section_num = 3 if data_stores else 2
    _add_heading_styled(doc, f"{section_num}. Passo a Passo do Processo",
                        level=1, color=TJMG_VINHO)

    if model.graph is None or len(model.graph.nodes) == 0:
        _add_styled_paragraph(doc, "Grafo do processo não disponível.",
                               font_size=10, italic=True, color=TJMG_CINZA)
    else:
        _generate_workflow_steps(
            doc, model, task_annotations, task_data_in,
            task_data_out, task_data_store
        )

    # ── FOOTER ──────────────────────────────────────────────────────
    doc.add_paragraph()
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line2.add_run("━" * 60)
    _style_run(run, font_size=8, color=TJMG_DOURADO)

    _add_styled_paragraph(
        doc,
        "Documento gerado automaticamente pelo Mapeador Inteligente — CEPROC/TJMG",
        font_size=8, color=TJMG_CINZA, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    # ── Serialize ───────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════
# WORKFLOW TRAVERSAL — BFS com narrativa
# ═══════════════════════════════════════════════════════════════════

def _generate_workflow_steps(doc, model, task_annotations,
                             task_data_in, task_data_out, task_data_store):
    """
    BFS traversal of the process graph, generating narrative text for each step.
    Handles system inheritance, data objects, annotations, and gateway branching.
    """
    G = model.graph
    if G is None:
        return

    # Find start events
    start_nodes = [
        eid for eid, elem in model.elements.items()
        if elem.element_type == "event" and "start" in elem.event_subtype
    ]
    if not start_nodes:
        # Fallback: nodes with in-degree 0
        start_nodes = [n for n in G.nodes() if G.in_degree(n) == 0]
    if not start_nodes:
        start_nodes = list(G.nodes())[:1]

    # BFS traversal
    visited = set()
    queue = deque(start_nodes)
    step_number = 0
    current_system = None  # System inheritance
    current_phase = None   # Phase/milestone tracking

    while queue:
        node_id = queue.popleft()
        if node_id in visited:
            continue
        visited.add(node_id)

        elem = model.elements.get(node_id)
        if elem is None:
            # Try next
            for succ in G.successors(node_id):
                if succ not in visited:
                    queue.append(succ)
            continue

        # ── Check for system update ─────────────────────────────
        if node_id in task_data_store:
            current_system = task_data_store[node_id]

        # ── Handle by element type ──────────────────────────────

        # Start event
        if elem.element_type == "event" and "start" in elem.event_subtype:
            event_name = elem.name or "Início"
            _add_styled_paragraph(
                doc,
                f"▶ Início do Processo: {event_name}",
                font_size=10, bold=True, color=TJMG_AZUL,
                space_before=12,
            )
            if elem.lane_name:
                _add_styled_paragraph(
                    doc, f"   Responsável inicial: {elem.lane_name}",
                    font_size=9, color=TJMG_CINZA,
                )

        # Intermediate event (milestone)
        elif elem.element_type == "event" and "intermediate" in elem.event_subtype:
            milestone_name = elem.name or "Marco"
            current_phase = milestone_name

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            run = p.add_run(f"━━━ MARCO: {milestone_name} ━━━")
            _style_run(run, font_size=11, color=TJMG_DOURADO, bold=True)

        # End event
        elif elem.element_type == "event" and "end" in elem.event_subtype:
            event_name = elem.name or "Fim"
            _add_styled_paragraph(
                doc,
                f"■ Fim do Processo: {event_name}",
                font_size=10, bold=True, color=TJMG_VINHO,
                space_before=12,
            )

        # Task/Activity
        elif elem.element_type == "task":
            step_number += 1
            lane_name = elem.lane_name or "—"
            task_name = elem.name or "Atividade sem nome"

            # Phase subtitle (if new)
            # Main step text
            system_text = f" no sistema {current_system}" if current_system else ""

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)

            run_num = p.add_run(f"Passo {step_number} — ")
            _style_run(run_num, font_size=10, color=TJMG_VINHO, bold=True)

            run_desc = p.add_run(f"{task_name}")
            _style_run(run_desc, font_size=10, bold=True)

            # Structured description
            _add_styled_paragraph(
                doc,
                f"O responsável {lane_name} realiza a atividade \"{task_name}\"{system_text}.",
                font_size=9, space_after=2,
            )

            # Insumos (inputs)
            inputs = task_data_in.get(node_id, [])
            if inputs:
                for inp in inputs:
                    _add_styled_paragraph(
                        doc, f"   📥 Documento Base/Insumo: {inp}",
                        font_size=9, color=TJMG_CINZA,
                    )

            # Produtos (outputs)
            outputs = task_data_out.get(node_id, [])
            if outputs:
                for out in outputs:
                    _add_styled_paragraph(
                        doc, f"   📤 Documento Gerado: {out}",
                        font_size=9, color=TJMG_CINZA,
                    )

            # Annotations (business rules)
            annotations = task_annotations.get(node_id, [])
            if annotations:
                for ann_text in annotations:
                    p_ann = doc.add_paragraph()
                    run_label = p_ann.add_run("   📝 Nota/Parâmetro: ")
                    _style_run(run_label, font_size=9, bold=True, color=TJMG_CINZA)
                    run_text = p_ann.add_run(ann_text)
                    _style_run(run_text, font_size=9, italic=True, color=TJMG_CINZA)

        # Gateway (decision)
        elif elem.element_type == "gateway":
            gate_name = elem.name or "Condição"
            successors = list(G.successors(node_id))

            if len(successors) > 1 and gate_name:
                _add_styled_paragraph(
                    doc,
                    f"🔀 A seguir, a condição avaliada é: {gate_name}",
                    font_size=10, bold=True, color=TJMG_AZUL,
                    space_before=8,
                )

                for succ_id in successors:
                    # Get edge label
                    edge_data = G.get_edge_data(node_id, succ_id, default={})
                    flow_name = edge_data.get("name", "")

                    succ_elem = model.elements.get(succ_id)
                    succ_name = succ_elem.name if succ_elem else succ_id

                    if flow_name:
                        _add_styled_paragraph(
                            doc,
                            f"     → Se \"{flow_name}\": seguir para \"{succ_name}\"",
                            font_size=9,
                        )
                    else:
                        _add_styled_paragraph(
                            doc,
                            f"     → Seguir para \"{succ_name}\"",
                            font_size=9,
                        )

        # ── Queue successors ────────────────────────────────────
        for succ in G.successors(node_id):
            if succ not in visited:
                queue.append(succ)
